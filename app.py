import os
import json
import time
import logging
import requests
import pdfplumber
import streamlit as st
import pandas as pd
import magic
import base64
import re
from io import BytesIO
from logging.handlers import RotatingFileHandler
from concurrent.futures import ThreadPoolExecutor
from functools import wraps
from dotenv import load_dotenv
from docx import Document
from PIL import Image
import pytesseract
import cv2
import numpy as np
import tempfile
import warnings
from pdfminer.pdfpage import PDFPage

# Suppress PDFPlumber warnings about CropBox
warnings.filterwarnings("ignore", category=UserWarning, message="CropBox missing from /Page")

# This will prevent the "CropBox missing from /Page" warnings from appearing in your logs

load_dotenv()

# Configure logging to a JSON file for parsing attempts
logging.basicConfig(
    filename='parsing_log.json',
    level=logging.INFO,
    format='%(message)s'
)

# Configure separate loggers for different types of output
# 1. JSON Logger for parsing attempts
json_logger = logging.getLogger('json_logger')
json_logger.setLevel(logging.INFO)
json_handler = RotatingFileHandler('parsing_data.json', maxBytes=10485760, backupCount=5)
json_logger.addHandler(json_handler)
json_logger.propagate = False  # Don't send to root logger

# 2. Regular logger for other messages
app_logger = logging.getLogger('app_logger')
app_logger.setLevel(logging.INFO)
app_handler = RotatingFileHandler('app.log', maxBytes=10485760, backupCount=5)
app_handler.setFormatter(logging.Formatter('%(asctime)s - %(levelname)s - %(message)s'))
app_logger.addHandler(app_handler)

# DeepSeek API configuration
DEEPSEEK_API_KEY = os.getenv("DEEPSEEK_API_KEY")
API_URL = "https://api.deepseek.com/v1/chat/completions"
HEADERS = {
    "Authorization": f"Bearer {DEEPSEEK_API_KEY}",
    "Content-Type": "application/json"
}

# Add this to the beginning of your file after the import statements

class JsonLogWriter:
    """Custom log handler that properly maintains valid JSON formatting"""
    
    def __init__(self, filename):
        self.filename = filename
        # Initialize the file with an empty array
        with open(self.filename, 'w') as f:
            f.write('[\n')
            f.write(']')
        self.entry_count = 0
    
    def write_entry(self, entry):
        """Write a new entry to the JSON log file while maintaining valid JSON structure"""
        with open(self.filename, 'r+') as f:
            # Read the file up to the last character (which should be ']')
            f.seek(0, 0)
            content = f.read()
            
            # Position cursor right before the closing bracket
            f.seek(len(content) - 1, 0)
            
            # Write a comma if this isn't the first entry
            if self.entry_count > 0:
                f.write(',\n')
            else:
                f.write('\n')
            
            # Write the new entry
            f.write(json.dumps(entry, indent=2))
            
            # Close the array again
            f.write('\n]')
            
            self.entry_count += 1
    
    def close(self):
        """No-op close method"""
        pass

# Create an instance of the JSON log writer
json_log = JsonLogWriter('parsing_log.json')


# Create a direct PDF display component using st.components.html
def pdf_viewer_html(pdf_data):
    """Display a PDF using HTML/JS viewer that works on first click."""
    base64_pdf = base64.b64encode(pdf_data).decode('utf-8')
    
    # Use PDF.js for reliable rendering
    pdf_display = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <script src="https://cdnjs.cloudflare.com/ajax/libs/pdf.js/2.11.338/pdf.min.js"></script>
        <style>
            #canvas_container {{
                width: 100%;
                height: 600px;
                overflow: auto;
                background: #333;
                text-align: center;
                border: 1px solid #666;
            }}
            #pdf-viewer {{
                width: 100%;
                border: none;
                height: 600px;
            }}
        </style>
    </head>
    <body>
        <div id="canvas_container">
            <iframe id="pdf-viewer" src="data:application/pdf;base64,{base64_pdf}#zoom=100"></iframe>
        </div>
    </body>
    </html>
    """
    
    # Use html component with a fixed height
    return st.components.v1.html(pdf_display, height=620)

# Rate limiting decorator
def rate_limit(max_calls=5, period=60):
    """
    Rate limit the API calls to avoid overloading the server.
    max_calls: Maximum number of calls allowed in the period
    period: Time period in seconds
    """
    calls = []
    
    def decorator(func):
        @wraps(func)
        def wrapper(*args, **kwargs):
            now = time.time()
            
            # Remove calls older than the period
            while calls and calls[0] < now - period:
                calls.pop(0)
            
            # If we've made too many calls, wait
            if len(calls) >= max_calls:
                wait_time = calls[0] + period - now
                if wait_time > 0:
                    st.info(f"Rate limit reached. Waiting {wait_time:.1f} seconds before next API call...")
                    time.sleep(wait_time)
            
            # Make the call and record the time
            result = func(*args, **kwargs)
            calls.append(time.time())
            return result
        return wrapper
    return decorator

# Improved retry decorator with exponential backoff and more attempts
def retry(max_attempts=5, initial_delay=2, backoff=2, max_timeout=60):
    def decorator_retry(func):
        def wrapper(*args, **kwargs):
            attempts = 0
            delay = initial_delay
            timeout = 30  # Start with default timeout
            
            while attempts < max_attempts:
                try:
                    # Increase timeout with each retry
                    if 'timeout' in kwargs:
                        kwargs['timeout'] = min(timeout, max_timeout)
                    
                    return func(*args, **kwargs)
                except Exception as e:
                    attempts += 1
                    if attempts == max_attempts:
                        app_logger.error(f"Maximum retry attempts ({max_attempts}) reached. Last error: {str(e)}")
                        raise
                    
                    logging.warning(f"API request failed (attempt {attempts}/{max_attempts}). Retrying in {delay}s...")
                    time.sleep(delay)
                    delay *= backoff
                    timeout *= 1.5  # Increase timeout for next attempt
            
            return None  # Should never reach here
        return wrapper
    return decorator_retry

# --- Extraction functions ---

def detect_file_type(file_bytes):
    """Detect the file type using python-magic."""
    file_type = magic.from_buffer(file_bytes.read(2048), mime=True)
    file_bytes.seek(0)
    return file_type

def extract_pdf_vector(uploaded_file):
    """Extract text using pdfplumber (vector/text layer) from a PDF."""
    try:
        with pdfplumber.open(uploaded_file) as pdf:
            text = "\n".join([page.extract_text() or "" for page in pdf.pages])
        return text.strip()
    except Exception as e:
        app_logger.error(f"Vector extraction error for {uploaded_file.name}: {str(e)}")
        return ""

def extract_pdf_ocr(uploaded_file):
    """Extract text using OCR from a PDF by converting pages to images."""
    try:
        text = ""
        with pdfplumber.open(uploaded_file) as pdf:
            for page in pdf.pages:
                im = page.to_image(resolution=300).original
                open_cv_image = cv2.cvtColor(np.array(im), cv2.COLOR_RGB2BGR)
                gray = cv2.cvtColor(open_cv_image, cv2.COLOR_BGR2GRAY)
                text += pytesseract.image_to_string(gray) + "\n"
        return text.strip()
    except Exception as e:
        app_logger.error(f"OCR extraction error for {uploaded_file.name}: {str(e)}")
        return ""

def extract_docx_text(uploaded_file):
    """Extract text from DOCX file with style retention."""
    try:
        doc = Document(uploaded_file)
        text = "\n".join([para.text for para in doc.paragraphs])
        return text.strip()
    except Exception as e:
        app_logger.error(f"DOCX extraction error for {uploaded_file.name}: {str(e)}")
        return ""

def extract_image_text(uploaded_file):
    """Extract text from image using Tesseract OCR with adaptive preprocessing."""
    try:
        image = Image.open(uploaded_file)
        open_cv_image = cv2.cvtColor(np.array(image), cv2.COLOR_RGB2BGR)
        gray = cv2.cvtColor(open_cv_image, cv2.COLOR_BGR2GRAY)
        processed = cv2.adaptiveThreshold(gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
                                          cv2.THRESH_BINARY, 11, 2)
        text = pytesseract.image_to_string(processed)
        return text.strip()
    except Exception as e:
        app_logger.error(f"Image OCR extraction error for {uploaded_file.name}: {str(e)}")
        return ""

def extract_text(uploaded_file):
    """Determine file type and extract text accordingly."""
    file_bytes = uploaded_file.getvalue()
    file_type = detect_file_type(uploaded_file)
    
    if 'pdf' in file_type:
        text = extract_pdf_vector(uploaded_file)
        if not text:
            text = extract_pdf_ocr(uploaded_file)
    elif 'msword' in file_type or 'officedocument' in file_type:
        text = extract_docx_text(uploaded_file)
    elif 'image' in file_type:
        text = extract_image_text(uploaded_file)
    else:
        st.warning(f"Unsupported file type ({file_type}). Falling back to metadata analysis.")
        text = uploaded_file.name  # fallback using filename as metadata
    return text

# --- Role-Specific Prompt Builder ---
def get_role_specific_prompt(role, job_description, resume_text, mode, include_contact_info=False):
    """
    Build a role-specific prompt for ATS resume evaluation.
    mode: "normal" or "deep"
    include_contact_info: Whether to ask DeepSeek to extract contact information
    """
    contact_info_request = """
Additionally, extract the mobile number and email address from the resume. 
Return these in the following format after the match percentage:
'Mobile: <extracted_mobile>' (Format as a simple number without spaces or symbols. If there's a country code like +1 or +91, extract only the 10-digit number. Examples: "+91 1234567890" → "1234567890")
'Email: <extracted_email>' (Standardize email format if needed, e.g., "name@domain.co" → "name@domain.com")
If you can't find a mobile number or email, respond with 'Mobile: N/A' or 'Email: N/A' respectively.
"""
    
    base_prompts = {
        "AI Specialist": f"""You are an advanced ATS scanner specialized in AI and Machine Learning recruitment. Evaluate the resume against the job description with precision and depth. For {mode.upper()} mode, assess the candidate's technical expertise in artificial intelligence, machine learning, deep learning, NLP, computer vision, algorithm development, and use of frameworks (e.g., TensorFlow, PyTorch). Consider innovation and problem-solving abilities.""",
        
        "Data Scientist": f"""You are an expert ATS scanner specialized in Data Science recruitment. Evaluate the resume against the job description with rigorous analysis. For {mode.upper()} mode, assess proficiency in statistical analysis, machine learning algorithms, data engineering, Python, SQL, data visualization (e.g., Tableau, Power BI), and data modeling.""",
        
        "Credit Research": f"""You are a top-tier ATS scanner specialized in Credit Research recruitment. Evaluate the resume against the job description with meticulous attention to International Credit research, buy side, distressed debt, stressed/distressed credit analysis, Investment grade, high yield, financial modeling, capital structure analysis, investment research, bond covenants, recommendation, research analyst. For {mode.upper()} mode, assess the candidate's expertise in International Credit research, buy side, distressed debt, stressed/distressed credit analysis, Investmestnt grade, high yield, financial modeling, capital structure analysis, investment research, bond covenants, recommendation, research analyst.""",
        
        "Equity Research": f"""You are a finance-focused ATS scanner specialized in Equity Research recruitment. Evaluate the resume against the job description with comprehensive scrutiny. For {mode.upper()} mode, assess the candidate's proficiency in financial modeling, valuation, earnings analysis, forecasting, fundamental and technical analysis, investment thesis, buy-side and sell-side research, industry analysis, market trends, and risk assessment, discounted cash flow (DCF), relative valuation, comparable company analysis (Comps), precedent transactions, enterprise value (EV), price-to-earnings (P/E) ratio, price-to-book (P/B) ratio, equity research reports, investment recommendations, earnings call summaries, company profiles, sector research, thematic reports, initiating coverage, and quarterly and annual reports, macroeconomic analysis, sectoral analysis, ESG (Environmental, Social, and Governance) research, growth stocks, value stocks, and momentum investing, Bloomberg, Reuters Eikon, FactSet, S&P Capital IQ,  Microsoft Excel, VBA, PowerPoint.""",
        
        "HR Specialist": f"""You are an advanced ATS scanner specialized in Human Resources recruitment. Evaluate the resume against the job description with detailed attention to HR competencies, such as talent acquisition, employee engagement, performance management, HR policy implementation, and regulatory compliance. For {mode.upper()} mode, assess Recruitment, Talent acquisition, Learning and development, Training need, Policy, assess strategic HR planning, communication skills, and proficiency in HRIS systems.""",
        
        "IT Support": f"""You are a proficient ATS scanner specialized in IT Support recruitment. Evaluate the resume against the job description with a focus on technical troubleshooting, customer support, hardware/software maintenance, and IT service management. For {mode.upper()} mode, assess experience in in O365, SharePoint, Azure, Outlook,  OneDrive, Teams, OS (Windows, Linux, MAC & Mobiles(android)),CCTV, Firewall, LAN, WAN, WIFI, Access Point, Printer, Video Conferencing Setup, VPN, Switching, Remote, Customer Handling Skill, Office Admin Role, Candidate Should have bachelor's degree (graduated after 2017) in IT, Certifications (CompTIA A+, ITIL) preferred, CCNA.""",
        
        "Presales Specialist": f"""You are a dynamic ATS scanner specialized in Presales Specialist recruitment. Evaluate the resume against the job description with emphasis on technical sales support, solution consulting, and client engagement. For {mode.upper()} mode, assess proficiency in technical presentations, product knowledge, and the ability to translate technical details into business benefits.""",
        
        "CRM & Lead Generation Specialist": f"""You are an expert ATS scanner specialized in CRM and Lead Generation recruitment. Evaluate the resume against the job description with a focus on customer relationship management, lead nurturing, and data-driven marketing strategies. For {mode.upper()} mode, assess proficiency in Fixed Income, Credit, Investment Grade, High Yield, Distressed, Leveraged Loans, Leveraged Finance, Special Situation ,Equities, Equities Sales,Equity, Equity Sales, Equity Broker, Equities Broker, Merger & Acquisition, Mergers & Acquisitions, M&A, Capital Market, DCM, ECM, Investment Banking, IB, Investment Research, Equity Capital Markets, Debt Capital Markets,Investment Management, Hedge Fund, Asset Management, Asset Manager, Brokerage, Broking, Sales & Trading.""",
        
        "IB Research": f"""You are a specialized ATS scanner focused on Investment Banking Research recruitment. Evaluate the resume against the job description with a detailed analysis of financial modeling, industry research, and market intelligence. For {mode.upper()} mode, assess the candidate's expertise in investment banking, ECM, DCM, M&A, merger, pitchbook, information memorandum, capital markets, comps, company profiles, teasers, industry research, valuation, and data or library services.""",
        
        "Business Research": f"""You are an expert ATS scanner specialized in Business Research recruitment. Evaluate the resume against the job description with comprehensive analysis of market research, competitive intelligence, and strategic planning. For {mode.upper()} mode, assess the candidate's proficiency in Market Research, Industry Analysis, Competitive Intelligence, Business Strategy, Financial Analysis, Data Interpretation, Trend Analysis, Feasibility Studies, Secondary Research, Primary Research, Forecasting, SWOT Analysis.""",
        
        "Digital Marketing": f"""You are an advanced ATS scanner specialized in Digital Marketing recruitment. Evaluate the resume against the job description with precision and depth. For {mode.upper()} mode, assess the candidate's expertise in Search Engine Optimization (SEO) including on-page and off-page strategies, organic traffic growth, content creation and management (blogs, articles, infographics, images, videos), and digital brand building. Evaluate experience in managing and optimizing paid advertising campaigns across Google Ads, LinkedIn Ads, and Facebook Ads, including ad copy creation, landing page optimization, performance tracking, ROI and KPI measurement. Assess proficiency in website management using WordPress, basic HTML knowledge, social media marketing across platforms such as LinkedIn, Facebook, Twitter, and coordination with third-party marketing agencies. Consider skills in lead generation, inbound marketing, branding, campaign reporting, communication abilities, ability to work independently under tight deadlines, and prior exposure to BFSI, media, advertising, PR, or publishing industries.""",
        
        "TEV Analyst": f"""You are an advanced ATS scanner specialized in Project Finance and Credit Research recruitment. Evaluate the resume against the job description with precision and depth. For {mode.upper()} mode, assess the candidate's expertise in Project Finance, Credit Appraisal, and Financial Advisory. Evaluate experience in preparation of Techno Economic Viability (TEV) Reports, Detailed Project Reports (DPRs), Lenders’ Independent Engineer (LIE) Reports, Debt Restructuring Studies, and Project Finance cashflow models for infrastructure and industrial projects. Assess proficiency in advanced Excel-based financial modeling, including project cash flows, debt structuring, sensitivity analysis, and financial viability assessment based on raw data and assumptions. Review experience in drafting Information Memoranda (IMs) and credit appraisal notes for banks, NBFCs, and investors. Evaluate ability to validate technical and financial assumptions using market intelligence, sectoral research, and industry benchmarking. Consider exposure to site visits, technical due diligence, project execution assessment (costs, timelines, risks), and coordination with stakeholders such as lenders, TEV/LIE consultants, valuers, legal advisors, and technical experts. Assess understanding of the Indian lending ecosystem including PSU banks, private sector banks, and NBFCs. Additionally, evaluate client-handling skills, relationship management with financial institutions and corporates, presentation capabilities during lender appraisals, ability to manage multiple projects independently under tight deadlines, regulatory compliance awareness, MIS reporting, and willingness to travel extensively across India."""

    }
    
    # Default prompt if no role is selected
    base_prompt = base_prompts.get(role, f"""You are an advanced document processing system with hybrid parsing capabilities.
Mode: {mode.upper()}
Analyze the given resume against the provided job description.
For normal mode, calculate the percentage match based on skills, experience, and keywords.
If in deep mode (activated when initial score is 0), cross-validate the extraction using 3 methods and select the best result.""")
    
    # Standard output format
    output_format = "Follow the format: 'Percentage Match: XX%'"
    
    if include_contact_info:
        output_format += contact_info_request
    else:
        output_format += " without additional commentary."
    
    # Complete prompt
    prompt = f"""{base_prompt}
{output_format}
Job Description:
{job_description}
Resume Content:
{resume_text}"""
    
    return prompt


# --- DeepSeek API call with retry and rate limit ---
@rate_limit(max_calls=5, period=60)
@retry(max_attempts=5, initial_delay=2, backoff=2, max_timeout=90)
def get_deepseek_response(prompt):
    """Get response from DeepSeek API with enhanced reliability."""
    # Trim the prompt if it's very long to reduce processing time
    if len(prompt) > 8000:
        prompt = prompt[:8000]
        
    payload = {
        "model": "deepseek-chat",
        "temperature": 0,
        "messages": [{"role": "user", "content": prompt}]
    }
    
    try:
        response = requests.post(API_URL, headers=HEADERS, json=payload, timeout=45)  # Increased timeout
        response.raise_for_status()
        return response.json()["choices"][0]["message"]["content"]
    except requests.exceptions.Timeout:
        # Special handling for timeout errors
        logging.warning("DeepSeek API timed out. Generating fallback response...")
        
        # Generate a fallback response to keep the application running
        if "match percentage" in prompt.lower():
            # Extract some keywords to make an educated guess at matching
            job_desc = prompt.split("Job Description:")[1].split("Resume Content:")[0] if "Job Description:" in prompt else ""
            resume = prompt.split("Resume Content:")[1] if "Resume Content:" in prompt else ""
            
            # Very basic fallback matching logic
            # Count matches of key terms
            if job_desc and resume:
                job_terms = set([word.lower() for word in job_desc.split() if len(word) > 4])
                resume_terms = set([word.lower() for word in resume.split() if len(word) > 4])
                overlap = job_terms.intersection(resume_terms)
                
                fallback_score = min(85, int(len(overlap) * 100 / (len(job_terms) or 1)))
                
                # Extract a possible email from the resume using regex
                email_match = re.search(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', resume)
                email = email_match.group(0) if email_match else "N/A"
                
                # Extract a possible phone number
                phone_match = re.search(r'(?:\+\d{1,3})?\s*\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}', resume)
                phone = re.sub(r'\D', '', phone_match.group(0))[-10:] if phone_match else "N/A"
                
                fallback_response = f"Percentage Match: {fallback_score}%\nMobile: {phone}\nEmail: {email}"
                return fallback_response
        
        # If all else fails, return a neutral response
        return "Percentage Match: 50%\nMobile: N/A\nEmail: N/A"
    except Exception as e:
        # Log error but don't expose full details to user
        app_logger.error(f"DeepSeek API error: {str(e)}")
        raise

def extract_percentage_and_contact(response_text):
    """Extract percentage match, mobile number, and email from API response with more robust patterns."""
    percentage = 0
    mobile = "N/A"
    email = "N/A"
    
    # Check if response is an error or fallback
    if response_text.startswith("ERROR:"):
        logging.warning(f"Cannot extract from error response: {response_text}")
        return 0, "N/A", "N/A"
    
    # Extract percentage - more flexible pattern
    try:
        # Try several possible formats for percentage
        patterns = [
            r'Percentage Match:\s*(\d+)%',
            r'Match:?\s*(\d+)%',
            r'Score:?\s*(\d+)%',
            r'(\d+)%\s*match',
            r'match\s*(\d+)%'
        ]
        
        for pattern in patterns:
            match = re.search(pattern, response_text, re.IGNORECASE)
            if match:
                percentage = int(match.group(1))
                break
        
        if percentage == 0:
            # Last resort - find any number followed by % symbol
            percent_match = re.search(r'(\d+)\s*%', response_text)
            if percent_match:
                percentage = int(percent_match.group(1))
    except Exception as e:
        app_logger.error(f"Error extracting percentage: {str(e)}")
        percentage = 0
    
    # Extract mobile number - more robust pattern
    try:
        # First try the exact format we asked for
        mobile_match = re.search(r'Mobile:\s*([0-9]+|N/A)', response_text)
        
        if mobile_match:
            mobile_text = mobile_match.group(1)
            if mobile_text != "N/A":
                # Extract the last 10 digits if the number is longer
                digits_only = re.sub(r'\D', '', mobile_text)
                if len(digits_only) >= 10:
                    mobile = digits_only[-10:]  # Get last 10 digits
                else:
                    mobile = digits_only  # Keep as is if less than 10 digits
            else:
                mobile = "N/A"
        else:
            # If exact format not found, try more flexible patterns
            phone_patterns = [
                r'(?:Mobile|Phone|Contact|Cell|Tel):\s*([+\d\s()\-.,]+)',
                r'(?:Mobile|Phone|Contact|Cell|Tel)[:\s]+([+\d\s()\-.,]+)',
                r'(?:Mobile|Phone|Contact|Cell|Tel)[:\s]*([+\d\s()\-.,]{8,20})'
            ]
            
            for pattern in phone_patterns:
                phone_match = re.search(pattern, response_text, re.IGNORECASE)
                if phone_match:
                    phone_text = phone_match.group(1).strip()
                    if phone_text and phone_text.lower() != "n/a":
                        # Clean and format the number
                        digits_only = re.sub(r'\D', '', phone_text)
                        if len(digits_only) >= 10:
                            mobile = digits_only[-10:]  # Get last 10 digits
                        elif len(digits_only) > 0:
                            mobile = digits_only
                    break
            
            # If still no match, look for any phone-like pattern in the response
            if mobile == "N/A":
                # Look for phone-like patterns anywhere in the text
                raw_phone_match = re.search(r'(?:\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}', response_text)
                if raw_phone_match:
                    phone_text = raw_phone_match.group(0)
                    digits_only = re.sub(r'\D', '', phone_text)
                    if len(digits_only) >= 10:
                        mobile = digits_only[-10:]  # Get last 10 digits
    except Exception as e:
        app_logger.error(f"Error extracting mobile: {str(e)}")
        mobile = "N/A"
        
    # Extract email - more robust pattern
    try:
        # Try the exact format we asked for
        email_match = re.search(r'Email:\s*([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}|N/A)', response_text)
        
        if email_match:
            email = email_match.group(1)
        else:
            # If exact format not found, try more flexible patterns
            email_patterns = [
                r'(?:Email|Mail|E-mail):\s*([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})',
                r'(?:Email|Mail|E-mail)[:\s]+([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})',
                r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'  # Find any email-like pattern
            ]
            
            for pattern in email_patterns:
                match = re.search(pattern, response_text, re.IGNORECASE)
                if match:
                    email_text = match.group(0) if pattern == email_patterns[-1] else match.group(1)
                    if email_text and email_text.lower() != "n/a":
                        email = email_text
                    break
    except Exception as e:
        app_logger.error(f"Error extracting email: {str(e)}")
        email = "N/A"
        
    return percentage, mobile, email

# Replace your process_resume function with this updated version
# that uses proper JSON logging

def process_resume(args):
    """
    Process a single document:
    1. First-pass analysis using primary extraction.
    2. If score == 0, perform deep-parsing mode with cross-validation.
    3. Extract contact information (mobile and email)
    Returns: (display filename, final percentage score, mobile, email, file object)
    """
    uploaded_file, job_description = args
    filename = uploaded_file.name
    rechecked = False
    try:
        # First-pass extraction
        resume_text = extract_text(uploaded_file)
        # Get selected role from session state (if none, default prompt will be used)
        role = st.session_state.get("selected_role", None)
        prompt_normal = get_role_specific_prompt(role, job_description, resume_text, mode="normal", include_contact_info=True)
        response_text = get_deepseek_response(prompt_normal)
        percentage, mobile, email = extract_percentage_and_contact(response_text)
        
        # Log normal attempt - using proper JSON logging
        log_entry = {
            "filename": filename,
            "mode": "normal",
            "score": percentage,
            "timestamp": time.time()
        }
        app_logger.info(f"Processed {filename} with score {percentage}%")  # Regular log
        json_log.write_entry(log_entry)  # JSON log
        
        # Recheck if initial score is 0
        if percentage == 0:
            rechecked = True
            if filename.lower().endswith(".pdf"):
                method_texts = {
                    'vector': extract_pdf_vector(uploaded_file),
                    'ocr': extract_pdf_ocr(uploaded_file),
                    'adaptive': extract_text(uploaded_file)
                }
                best_text = max(method_texts.values(), key=lambda t: len(t))
                prompt_deep = get_role_specific_prompt(role, job_description, best_text, mode="deep", include_contact_info=True)
                response_text = get_deepseek_response(prompt_deep)
                percentage, mobile, email = extract_percentage_and_contact(response_text)
            else:
                prompt_deep = get_role_specific_prompt(role, job_description, resume_text, mode="deep", include_contact_info=True)
                response_text = get_deepseek_response(prompt_deep)
                percentage, mobile, email = extract_percentage_and_contact(response_text)
            
            # Log deep-parsing attempt - using proper JSON logging
            log_entry = {
                "filename": filename,
                "mode": "deep",
                "score": percentage,
                "timestamp": time.time()
            }
            app_logger.info(f"Reprocessed {filename} in deep mode with score {percentage}%")  # Regular log
            json_log.write_entry(log_entry)  # JSON log
        
        display_name = filename + (" [†]" if rechecked else "")
        return (display_name, percentage, mobile, email, uploaded_file)
    except Exception as e:
        app_logger.error(f"Error processing {filename}: {str(e)}")  # Regular log
        
        # Log error - using proper JSON logging
        log_entry = {
            "filename": filename,
            "mode": "error",
            "error": str(e),
            "timestamp": time.time()
        }
        json_log.write_entry(log_entry)  # JSON log
        
        return (filename, 0, "N/A", "N/A", uploaded_file)

def process_documents_with_rate_limit(uploaded_files, job_description):
    """Process documents with rate limiting to avoid overwhelming the API."""
    results = []
    
    with st.spinner("Processing documents..."):
        # Reduce max_workers to avoid overwhelming the API
        with ThreadPoolExecutor(max_workers=2) as executor:  # Reduced from 5 to 2
            args = [(file, job_description) for file in uploaded_files]
            
            # Process in batches to avoid too many concurrent requests
            batch_size = 2
            for i in range(0, len(args), batch_size):
                batch_args = args[i:i+batch_size]
                batch_results = list(executor.map(process_resume, batch_args))
                results.extend(batch_results)
                
                # Add a small delay between batches
                if i + batch_size < len(args):
                    time.sleep(2)
    
    return results

def copy_file_to_folder(file_obj, filename, target_folder):
    """Copy the file from memory to the target folder."""
    if not os.path.exists(target_folder):
        os.makedirs(target_folder)
    # Write file content to disk
    with open(os.path.join(target_folder, filename), 'wb') as f:
        f.write(file_obj.getvalue())

def display_pdf(file_obj):
    """Create a base64 encoded version of the PDF for embedding in HTML."""
    # Pre-load the PDF data to cache it
    pdf_data = file_obj.getvalue()
    base64_pdf = base64.b64encode(pdf_data).decode('utf-8')
    
    # Set iframe attributes for better viewing experience
    pdf_display = f'''
    <iframe 
        src="data:application/pdf;base64,{base64_pdf}" 
        width="100%" 
        height="600" 
        type="application/pdf"
        style="border: none; display: block; min-height: 500px;"
        allow="fullscreen" 
    ></iframe>
    '''
    return pdf_display

def display_docx(file_obj):
    """Display DOCX file by converting to HTML."""
    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
        tmp.write(file_obj.getvalue())
        tmp_path = tmp.name
    
    try:
        # Create an HTML version for display
        doc = Document(tmp_path)
        html_content = "<div style='width:100%; height:600px; overflow:auto; border:1px solid #ccc; padding:10px;'>"
        for para in doc.paragraphs:
            html_content += f"<p>{para.text}</p>"
        html_content += "</div>"
        
        # Clean up the temp file
        os.unlink(tmp_path)
        return html_content
    except Exception as e:
        os.unlink(tmp_path)
        return f"<div class='error'>Error displaying document: {str(e)}</div>"

def display_image(file_obj):
    """Display an image in HTML."""
    file_type = detect_file_type(file_obj)
    base64_img = base64.b64encode(file_obj.getvalue()).decode('utf-8')
    img_display = f'<img src="data:{file_type};base64,{base64_img}" style="max-width:100%; max-height:600px;">'
    return img_display

# Another alternative - use a direct approach
def show_pdf_direct(file_obj):
    """Show PDF directly using a simple approach"""
    # Get the PDF data and encode it
    pdf_bytes = file_obj.getvalue()
    base64_pdf = base64.b64encode(pdf_bytes).decode('utf-8')
    
    # Create an iframe with base64 encoded data (simpler approach)
    pdf_html = f"""
    <iframe 
        src="data:application/pdf;base64,{base64_pdf}" 
        width="100%" 
        height="600px" 
        style="border:none;" 
        allowfullscreen="true">
    </iframe>
    """
    
    # Render the HTML with unsafe_allow_html=True
    st.markdown(pdf_html, unsafe_allow_html=True)
    
    # Also provide a download button
    st.download_button(
        "Download PDF",
        file_obj.getvalue(),
        file_name=file_obj.name,
        mime="application/pdf"
    )

# Replace your PDF display logic with one of these approaches
def display_resume(file_obj):
    """Display a resume based on its file type."""
    file_type = detect_file_type(file_obj)
    
    if 'pdf' in file_type:
        # Try multiple approaches until one works
        try:
            # Option 1: Direct PDF display (simplest)
            show_pdf_direct(file_obj)
        except Exception as e:
            st.error(f"Error displaying PDF: {str(e)}")
            st.download_button(
                "Download PDF instead", 
                file_obj.getvalue(), 
                file_name=file_obj.name,
                mime="application/pdf"
            )
    elif 'msword' in file_type or 'officedocument' in file_type:
        html_content = display_docx(file_obj)
        st.markdown(html_content, unsafe_allow_html=True)
    elif 'image' in file_type:
        html_content = display_image(file_obj)
        st.markdown(html_content, unsafe_allow_html=True)
    else:
        st.error(f"Unsupported file type for preview: {file_type}")
        
    return None  # Return None since we're displaying directly

# --- Streamlit UI ---
st.set_page_config(page_title="CVPerfect - Resume Matching", layout="wide")

# Set custom CSS to improve PDF preview
st.markdown("""
<style>
    .pdf-container iframe {
        width: 100%;
        height: 600px;
        display: block;
    }
    .stDataFrame tbody tr:hover {
        background-color: rgba(0, 0, 255, 0.1);
        cursor: pointer;
    }
    .file-selector {
        margin-bottom: 10px;
    }
    .preview-container {
        border: 1px solid #ddd;
        border-radius: 5px;
        padding: 10px;
        height: 600px;
        overflow: auto;
    }
</style>
""", unsafe_allow_html=True)

st.header("CVPerfect - An AI-Powered CV Matching Engine")
st.subheader("Find the Perfect Fit—Every Time.")  # Added tagline

# Job description input
job_description = st.text_area("Paste Job Description:", height=200, key="input")

# --- Role Selection Buttons ---
st.markdown("### Select Role for Resume Evaluation")
if "selected_role" not in st.session_state:
    st.session_state.selected_role = None

col1, col2, col3, col4, col5, col6, col7, col8, col9, col10, col11 = st.columns(11)
if col1.button("AI Specialist"):
    st.session_state.selected_role = "AI Specialist"
if col2.button("Data Scientist"):
    st.session_state.selected_role = "Data Scientist"
if col3.button("Credit Research"):
    st.session_state.selected_role = "Credit Research"
if col4.button("Equity Research"):
    st.session_state.selected_role = "Equity Research"
if col5.button("HR Specialist"):
    st.session_state.selected_role = "HR Specialist"
if col6.button("IT Support"):
    st.session_state.selected_role = "IT Support"
if col7.button("Presales Specialist"):
    st.session_state.selected_role = "Presales Specialist"
if col8.button("CRM & Lead Gen"):
    st.session_state.selected_role = "CRM & Lead Generation Specialist"
if col9.button("IB Research"):
    st.session_state.selected_role = "IB Research"
if col10.button("Business Research"):
    st.session_state.selected_role = "Business Research"
if col10.button("Digital Marketing"):
    st.session_state.selected_role = "Digital Marketing"
if col11.button("TEV Analyst"):
    st.session_state.selected_role = "TEV Analyst"

if st.session_state.selected_role:
    st.info(f"Selected Role: {st.session_state.selected_role}")
else:
    st.warning("⚠️ Please select a job role before proceeding.")
    st.stop()  # This will prevent further execution if no role is selected
    
# File uploader for multiple documents (PDF, DOCX, Images)
uploaded_files = st.file_uploader(
    "Upload Documents (PDF, DOCX, Images)", 
    type=["pdf", "docx", "png", "jpg", "jpeg"],
    accept_multiple_files=True,
    key = "file_upload"
)

# Store processed results in session state so they persist on widget interactions
if "processed_results" not in st.session_state:
    st.session_state.processed_results = None

# Store the currently selected file for preview
if "selected_file_for_preview" not in st.session_state:
    st.session_state.selected_file_for_preview = None

# Flag to trigger rerun when file is selected
if "file_selected" not in st.session_state:
    st.session_state.file_selected = False

# Store the preview HTML to avoid recomputation
if "preview_html" not in st.session_state:
    st.session_state.preview_html = None

if st.button("Process Documents") and uploaded_files:
    if not job_description.strip():
        st.warning("Please enter a job description")
        st.stop()
    
    # Use the rate-limited processing function
    results = process_documents_with_rate_limit(uploaded_files, job_description)
    st.session_state.processed_results = results
    
# Only display filtering and copying if we have processed results in session state
if st.session_state.processed_results:
    processed_list = []
    for res in st.session_state.processed_results:
        processed_list.append({
            "Filename": res[0],
            "Match Percentage": res[1],
            "Mobile Number": res[2],
            "Email ID": res[3],
            "FileObj": res[4]
        })
        
    # Display file counts below the slider
    st.write(f"Number of uploaded files: {len(processed_list)}")
    
    # Create dataframe for display
    df = pd.DataFrame(processed_list)[["Filename", "Match Percentage", "Mobile Number", "Email ID"]]
    df.index = df.index + 1  # Start ranking from 1
    
    # --- Slider for filtering based on percentage match ---
    
    # Slider with range selection from 0 to 100, defaulting to (0, 100)
    min_score, max_score = st.slider("Filter by Percentage Match", 0, 100, (0, 100))
    st.markdown(
        """
        <style>
        .tick-labels {
            display: flex;
            justify-content: space-between;
            margin-top: -10px;
        }
        </style>
        <div class="tick-labels">
            <span>0</span>
            <span>25</span>
            <span>50</span>
            <span>75</span>
            <span>100</span>
        </div>
        """,
        unsafe_allow_html=True
    )

    st.write(f"Selected Range: {min_score} to {max_score}")
    
    # Filter the processed list based on the selected range
    filtered = [item for item in processed_list if min_score <= item["Match Percentage"] <= max_score]
    st.write(f"Number of filtered files: {len(filtered)}")
    
    if len(filtered) == 0:
        st.warning("No file found above the selected percentage match.")
        filtered_df = pd.DataFrame(columns=["Filename", "Match Percentage", "Mobile Number", "Email ID"])
    else:
        filtered_df = pd.DataFrame(filtered)[["Filename", "Match Percentage", "Mobile Number", "Email ID"]]
        filtered_df.index = filtered_df.index + 1  # Start ranking from 1
    
    # Create two columns for split view (table and preview)
    
    col1, col2 = st.columns([2, 3])
    
    # Use this in the main UI section like this:
    # pdf_viewer_html(st.session_state.selected_file_for_preview.getvalue())
    
    with col1:
        # Display filtered results
        st.subheader("Filtered Ranking Results")
        st.write("Select a file number to preview the resume")
        
        # Create a dropdown for file selection
        file_options = [f"{i+1}. {item['Filename']}" for i, item in enumerate(filtered)]
        
        if file_options:
            selected_option = st.selectbox(
                "Choose a resume to preview:",
                options=file_options,
                key="file_selector"
            )
            
            # Extract the index from the selected option
            selected_index = int(selected_option.split('.')[0]) - 1
            
            # Preload the file object for preview
            file_obj = filtered[selected_index]["FileObj"]
            
            # Pre-process the file for preview and store in session state
            # This happens before any UI updates, ensuring preview is ready
            if "current_preview_id" not in st.session_state or st.session_state.current_preview_id != id(file_obj):
                with st.spinner("Preparing document preview..."):
                    file_type = detect_file_type(file_obj)
                    if 'pdf' in file_type:
                        preview_html = display_pdf(file_obj)
                    elif 'msword' in file_type or 'officedocument' in file_type:
                        preview_html = display_docx(file_obj)
                    elif 'image' in file_type:
                        preview_html = display_image(file_obj)
                    else:
                        preview_html = f"<div>Unsupported file type: {file_type}</div>"
                    
                    # Store both the file object and the rendered HTML
                    st.session_state.selected_file_for_preview = file_obj
                    st.session_state.preview_html = preview_html
                    st.session_state.current_preview_id = id(file_obj)
            
            # Display the dataframe
            st.dataframe(
                filtered_df,
                use_container_width=True,
                column_config={
                    "Match Percentage": st.column_config.ProgressColumn(
                        format="%d%%",
                        min_value=0,
                        max_value=100,
                    ),
                    "Filename": st.column_config.TextColumn(
                        "Filename",
                    )
                }
            )
        
        # Download and Copy buttons
        csv = filtered_df.to_csv(index=False).encode('utf-8')
        target_folder = "filtered_output"

        b1, b2 = st.columns([1, 1])
        with b1:
            st.download_button(
                label="Download Results as CSV",
                data=csv,
                file_name="filtered_document_ranking.csv",
                mime="text/csv"
            )

        with b2:
            if st.button("Copy Files to Folder", key="copy_button"):
                for item in filtered:
                    original_name = item["Filename"].replace(" [†]", "")
                    copy_file_to_folder(item["FileObj"], original_name, target_folder)
                st.success(f"Copied {len(filtered)} file(s) to folder '{target_folder}'.")
    
    with col2:
        st.subheader("Resume Preview")
        
        # Display the pre-processed HTML immediately without recalculation
        if st.session_state.preview_html:
            preview_container = st.container()
            with preview_container:
                st.markdown(
                    f'<div class="pdf-container">{st.session_state.preview_html}</div>',
                    unsafe_allow_html=True
                )
        else:
            st.info("Select a file to preview")


elif uploaded_files and not job_description:
    st.warning("Please enter a job description before processing")